import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout(doc, title, text, bg_hex="F0FDF4", border_hex="0F766E"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(15, 118, 110)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def create_document():
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    
    # Colors
    TEAL = RGBColor(15, 118, 110)
    EMERALD = RGBColor(5, 150, 105)
    DARK = RGBColor(15, 23, 42)
    GRAY = RGBColor(71, 85, 105)
    GOLD = RGBColor(217, 119, 6)
    
    # ==================== TITLE COVER / HEADER ====================
    p_pre = doc.add_paragraph()
    p_pre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pre.paragraph_format.space_before = Pt(10)
    p_pre.paragraph_format.space_after = Pt(2)
    r_uni = p_pre.add_run("QO‘QON UNIVERSITETI")
    r_uni.font.name = 'Calibri'
    r_uni.font.size = Pt(12)
    r_uni.bold = True
    r_uni.font.color.rgb = GOLD
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("«MUQIMIY AQL MARKAZI»\nINTERAKTIV AI KIOSK PLATFORMASI")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(22)
    r_title.bold = True
    r_title.font.color.rgb = TEAL
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(18)
    r_sub = p_sub.add_run("Tizim Arxitekturasi, Multi-LLM Intellekti, Ovoz Sintezi (TTS/STT)\nva Mohira AI Avatarining Ishlash Mexanizmi bo‘yicha To‘liq Texnik Hujjat")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = GRAY
    r_sub.italic = True
    
    # Meta table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Loyiha nomi:", "Muqimiy Aql Markazi — Smart AI Kiosk"),
        ("Tashkilot:", "Qo'qon Universiteti"),
        ("Asosiy Texnologiyalar:", "Next.js 14, React 18, TypeScript, TailwindCSS, ElevenLabs TTS, Groq Cloud LLM"),
        ("Hujjat Versiyasi:", "v2.4 (To'liq arxitektura va Avatar qo'llanmasi)")
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.name = 'Calibri'
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = TEAL
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = DARK
    
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    
    # Helpers
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = TEAL
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = EMERALD
        return h

    def add_body_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.bold = True
            rb.font.name = 'Calibri'
            rb.font.size = Pt(10.5)
            rb.font.color.rgb = DARK
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.italic = italic
        r.font.color.rgb = DARK
        return p

    # Chapter 1
    add_heading_1("1. LOYIHANING MAQSADI VA KONSEPSIYASI")
    add_body_p("«Muqimiy Aql Markazi» — Qo'qon Universiteti qoshida tashkil etilgan bo'lib, yoshlar, talabalar va keng jamoatchilikka zamonaviy bilimlar, xorijiy tillar, raqamli kasblar, qonuniy migratsiya hamda psixologik ko'mak berishga mo'ljallangan interaktiv Kiosk tizimidir.")
    add_body_p("Tizimning bosh virtual yo'lboshchisi sifatida milliy libosdagi o'zbek qizi — Mohira AI tanlangan. Foydalanuvchi sensorli ekran orqali yoki to'g'ridan-to'g'ri tabiiy ovozli muloqot orqali platformaning barcha imkoniyatlaridan foydalana oladi.")
    
    add_callout(doc, "Platformaning Asosiy Xususiyatlari", 
        "• 100% O'zbek tilida tabiiy jonli ovozli muloqot (ElevenLabs Neural Voice).\n"
        "• 3 ta videodan tashkil topgan Ultra-Realistik Video Avatar (Mohira).\n"
        "• Uzluksiz Multi-LLM zaxira arxitekturasi (Groq -> DeepSeek -> Gemini -> KieAI).\n"
        "• 5 ta mustaqil ta'lim va ko'mak yo'nalishlari integratsiyasi.\n"
        "• Psixologik ko'mak bo'limida 15 ta Solfeggio va Handpan yuqori Hz chastotali miyani tozalovchi shifobaxsh musiqalar.")

    # Chapter 2
    add_heading_1("2. UMUMIY TIZIM ARXITEKTURASI VA QATLAMLAR")
    add_body_p("Platforma modulli va qatlamli (Multi-tier) arxitekturaga ega bo'lib, har bir qatlam o'zining aniq vazifasini bajaradi:")
    
    # Table of Layers
    arch_table = doc.add_table(rows=6, cols=3)
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Qatlam", "Texnologiya", "Asosiy Vazifasi"]
    for i, h in enumerate(headers):
        cell = arch_table.rows[0].cells[i]
        set_cell_background(cell, "0F766E")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, 80, 80, 100, 100)
    
    layers_data = [
        ("Frontend / UI Qatlami", "Next.js 14, React 18, TailwindCSS, Framer Motion", "Kiosk interfeysi, 560px x 670px Video Avatar, vertikal svayper va sensorli boshqaruv"),
        ("Ovoz Qabul Qilish (STT)", "Web Speech Recognition API", "Foydalanuvchining jonli nutqini real vaqtda matnga aylantirish"),
        ("Server / API Qatlami", "Next.js Route Handlers (/api/chat, /api/tts)", "Prompt muhandisligi, AI modellararo zaxira marshrutizatsiyasi, ovoz sintezi keshlanishi"),
        ("Sun'iy Intellekt (LLM)", "Groq (LLaMA 3.3 / GPT-OSS-120b), DeepSeek, Gemini", "Kontekstli, samimiy va o'zbekona tushuntiruvchi aqlli javoblar generatsiyasi"),
        ("Ovoz Sintezi (TTS)", "ElevenLabs Multilingual v2, Microsoft Edge Neural TTS", "Matnni jonli inson ovoziga teng bo'lgan MP3 oqimiga aylantirish")
    ]
    
    for row_idx, data in enumerate(layers_data, start=1):
        for col_idx, val in enumerate(data):
            cell = arch_table.rows[row_idx].cells[col_idx]
            bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 70, 70, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = TEAL

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Chapter 3: Avatar
    add_heading_1("3. MOHIRA AI AVATARINING ISHLASH MEXANIZMI (CHUQUROQ TAHLIL)")
    add_body_p("Mohira avatari platformadagi eng muhim interaktiv element hisoblanadi. U `RealisticVideoAvatar.tsx` komponenti orqali boshqariladi va quyidagi 3 ta haqiqiy video aktivdan foydalanadi:")
    
    # 3 Videos table
    v_table = doc.add_table(rows=4, cols=3)
    v_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Video Fayl", "Hajmi & Davomiyligi", "Vazifasi va Vizual Ko'rinishi"]):
        cell = v_table.rows[0].cells[i]
        set_cell_background(cell, "059669")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, 80, 80, 100, 100)
    
    v_data = [
        ("public/avatar/greeting.mp4", "4.6 MB (4-5 sek)", "O'zbekona mehmondo'stlik ta'zimi — qo'lini ko'ksiga qo'yib, jilmayib foydalanuvchini qutlaydi."),
        ("public/avatar/idle.mp4", "5.4 MB (Tsiklik - loop)", "Kutish holati — ko'z pirpiratish, yengil tabassum va sokin nafas olish. Foydalanuvchi gapirayotganda va jim turganda ishlaydi."),
        ("public/avatar/talking.mp4", "5.0 MB (Tsiklik - loop)", "Gapirish holati — lablarning tabiiy nutq artikulyatsiyasi, bosh va mimikalarning so'zlashuv ritmiga mos harakati.")
    ]
    for row_idx, data in enumerate(v_data, start=1):
        for col_idx, val in enumerate(data):
            cell = v_table.rows[row_idx].cells[col_idx]
            bg = "F0FDF4" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 70, 70, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK
            if col_idx == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2("Avatar Holatlar Mashinasi (State Machine) va Ovoz Sinxronizatsiyasi:")
    add_body_p("Avatar holati (State) va audio pleyer o'rtasida to'liq sinxron bog'liqlik o'rnatilgan:")
    add_body_p("1. greeting (Qutlash): Sahifa birinchi marta yuklanganda Mohira qo'lini ko'ksiga qo'yadi va salomlashish audiosi ijro etiladi. Audio tugashi bilan avtomatik `onGreetingEnd()` chaqirilib, holat `idle`ga o'tadi.", bold_prefix="• ")
    add_body_p("2. idle (Sokin kutish): Mohira jilmayib turadi, `idle.mp4` fonda uzluksiz o'ynaydi.", bold_prefix="• ")
    add_body_p("3. listening (Eshitish): Foydalanuvchi gapira boshlaganda avatar chetlarida qizil-pushti nur porlaydi va «Sizni eshitmoqdaman, gapiring...» nishoni chiqadi.", bold_prefix="• ")
    add_body_p("4. thinking (O'ylash): Ovoz yozib olinib, AI modelga so'rov ketganda avatar atrofida binafsharang nur paydo bo'lib, intellektual jarayon kechayotganini ko'rsatadi.", bold_prefix="• ")
    add_body_p("5. speaking (Gapirish): AI javob audiosi kelishi bilan audio pleyer `play()` bo'ladi. Shu onning o'zida `talking.mp4` ning `opacity`si 100% bo'lib, `idle.mp4` to'xtatiladi. Audio tugagan soniyada `ended` hodisasi orqali `idle.mp4`ga silliq (smooth crossfade) qaytadi.", bold_prefix="• ")

    add_heading_2("Dinamik Nur va Vizual Aura (Glow System):")
    add_body_p("Kiosk monitorida hashamatli ko'rinish berish uchun avatarning orqasida CSS `blur-xl animate-pulse` filtrlari qo'llangan:")
    add_body_p("• Gapirayotganda: Yashil-zumrad nur (Emerald #059669 / Teal #0F766E)\n"
               "• Eshitayotganda: Qizil-pushti nur (Rose #E11D48 / Pink #F43F5E)\n"
               "• O'ylanayotganda: Binafsharang nur (Purple #9333EA)\n"
               "• Salomlashayotganda: Oltin-sariq nur (Amber #D97706 / Gold)")

    # Chapter 4
    add_heading_1("4. OVOZ VA MULTI-LLM INTIELLEKTI ARXITEKTURASI")
    add_heading_2("A. Zaxirali Multi-LLM Tizimi (/api/chat)")
    add_body_p("Kioskning uzluksiz 24/7 ishlashini ta'minlash uchun tizim 4 pog'onali avtomatik kaskadli zaxiraga ega:")
    add_body_p("1. Groq Cloud (openai/gpt-oss-120b yoki LLaMA-3.3-70b) — Asosiy model. Javob qaytarish tezligi 200-400ms ni tashkil etadi.", bold_prefix="1-Pog'ona: ")
    add_body_p("2. DeepSeek Chat V3 — 1-zaxira. Agar Groq API band bo'lsa yoki limit to'lsa, tizim sezilarsiz tarzda DeepSeek'ga ulanadi.", bold_prefix="2-Pog'ona: ")
    add_body_p("3. Google Gemini 1.5 Flash — 2-zaxira. Kuchli tahliliy model.", bold_prefix="3-Pog'ona: ")
    add_body_p("4. KIE AI / Mahalliy qoidalar — Favqulodda offline holatlar uchun zaxira.", bold_prefix="4-Pog'ona: ")

    add_heading_2("B. Neyron Ovoz Sintezi (/api/tts)")
    add_body_p("Matndan nutq hosil qilishda ikkita zamonaviy neyron tarmoq ishlatiladi:")
    add_body_p("• ElevenLabs Multilingual v2 (Model: EXAVITQu4vr4xnSDxMaL - Bella ovozi): Eng yuqori darajadagi tabiiy o'zbekcha ayol ovozi. Nafas olish, his-tuyg'u va ravon intonatsiyaga ega.\n"
               "• Microsoft Edge Neural TTS (uz-UZ-MadinaNeural va uz-UZ-SardorNeural): 100% bepul va cheklovsiz zaxira ovoz generatori.\n"
               "• Tarixiy shaxslar uchun: Amir Temur (Adam - salobatli sulton), Alisher Navoiy (George - shoirona dono), Mirzo Ulug'bek (Daniel - intellektual olim) kabi maxsus tematik ovozlar sozlangan.")

    # Chapter 5
    add_heading_1("5. ASOSIY 5 TA YO‘NALISH VA MODULLAR TAVSIFI")
    
    modules = [
        ("1. Boshqaruv Paneli (Kiosk Dashboard)", "Asosiy kirish sahifasi. 5 ta yo'nalishni vertikal animatsiyali svayper orqali taqdim etadi. Chap tomonda 560px x 670px o'lchamli to'liq ekranli Mohira avatari turadi. Foydalanuvchi bilan hands-free ovozli muloqot qiladi."),
        ("2. Chet Tillari Hubi (/language)", "Ibrat Farzandlari platformasi bilan integratsiya. Ingliz, Rus, Xitoy, Nemis, Ispan va Turk tillarini bepul o'rganish yo'nalishlari. Mohira har bir til bo'yicha yo'l xaritasini ovozli tushuntirib beradi."),
        ("3. Zamonaviy Kasblar Hubi (/career)", "Ustoz AI bilan hamkorlik. Web dasturlash, Grafik dizayn, SMM, Video montaj, 3D modellashtirish va Sun'iy intellekt muhandisligi yo'nalishlari bo'yicha kasbiy tavsiyalar."),
        ("4. Xavfsiz Migratsiya Hubi (/migration)", "Xorijda qonuniy ta'lim va ishlash dasturlari (Germaniya Ausbildung, Janubiy Koreya EPS E-9, Yaponiya Tokutei Ginou, Buyuk Britaniya, Polsha). Noqonuniy migratsiya xavflaridan ogohlantiruvchi qonuniy yo'riqnomalar."),
        ("5. Psixologik Ko'mak & 15 ta Hz Shifobaxsh Musiqalar (/psychology)", "3 savolli tezkor stress va kayfiyat diagnostika testi. Diagnostika natijasiga ko'ra 15 ta maxsus Solfeggio va Handpan yuqori chastotali (1111 Hz, 963 Hz, 852 Hz, 528 Hz, 432 Hz) miyani tozalovchi va tinchlantiruvchi musiqalar avtomatik tavsiya etiladi."),
        ("6. Virtual Tarixiy Allomalar Muzeyi (/history)", "Amir Temur, Alisher Navoiy, Mirzo Ulug'bek, Mahmudxo'ja Behbudiy va Abdulla Avloniy bilan sun'iy intellekt orqali jonli muloqot. Har bir allomaning o'z ovozi va shaxsiy falsafiy bilimlari kiritilgan.")
    ]
    
    for title, desc in modules:
        add_body_p(desc, bold_prefix=f"{title}: ")

    # Chapter 6: Hz music details
    add_heading_1("6. PSIXOLOGIYA BO‘LIMIDAGI 15 TA SHIFOBAXSH HZ MUSIQALAR")
    add_body_p("Foydalanuvchilarning ruhiy holatini yaxshilash va miyadagi ortiqcha xavotirni yuvish uchun 15 ta ilmiy isbotlangan Solfeggio va Handpan chastotalari joylashtirilgan:")
    
    hz_table = doc.add_table(rows=6, cols=3)
    hz_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Kategoriya", "Chastota (Hz) & Ijrochi", "Miyaga va Ruhiyatga Ta'siri"]):
        cell = hz_table.rows[0].cells[i]
        set_cell_background(cell, "D97706")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, 80, 80, 100, 100)
    
    hz_data = [
        ("Yuqori Stress & Xavotir", "1111 Hz Healing Handpan (Malte Marten & Lynxk)\n963 Hz Pineal Gland & Pure Brain Reset\n852 Hz Overthinking & Anti-Anxiety Flow", "Miyadagi barcha xavotir va chuqur stressni tozalash, ortiqcha o'y-xayollarni o'chirish"),
        ("Charchoq & Bosim", "528 Hz Seeds of Calm Handpan (Malte Marten)\n741 Hz Mental Detox\n396 Hz Release Pressure", "Mo'jizaviy chastota. Ruhiy quvvatni tiklaydi, aqliy toksinlarni haydaydi"),
        ("Tinchlik & Meditatsiya", "432 Hz Golden Light of Peace\n6 Hz Theta Wave Deep Mind Reset\n432 Hz Crystal Singing Bowls", "Miyaning chuqur neyronlarini yangilovchi Teta to'lqinlari, mayin xotirjamlik"),
        ("Qalb Jarohatlari & Dalda", "639 Hz Heart Chakra & Emotional Healing\n174 Hz Natural Pain Relief", "Mehr va iliqlik bag'ishlovchi ohang, ruhiy zo'riqishni yengillashtirish"),
        ("To'liq Shifo Sayohati", "9 Solfeggio Frequencies Full Journey\n432 Hz Yashil O'rmon & Handpan", "Barcha 9 ta shifobaxsh chastotaning to'liq tiklanish kaskadi")
    ]
    
    for row_idx, data in enumerate(hz_data, start=1):
        for col_idx, val in enumerate(data):
            cell = hz_table.rows[row_idx].cells[col_idx]
            bg = "FFFBEB" if row_idx % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 70, 70, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = GOLD

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Chapter 7: File structure
    add_heading_1("7. LOYIHANING ASOSIY FAYLLAR XARITASI")
    add_body_p("• `src/components/avatar/RealisticVideoAvatar.tsx` — 3 videoni boshqaruvchi, CSS aura va nishonlarga ega asosiy Avatar komponenti.\n"
               "• `src/app/api/chat/route.ts` — Groq, DeepSeek va Gemini kaskadli Multi-LLM routeri.\n"
               "• `src/app/api/tts/route.ts` — ElevenLabs va Microsoft Edge TTS neyron ovoz sintezatori.\n"
               "• `src/app/dashboard/page.tsx` — 560px x 670px o'lchamli Avatar va 5 yo'nalishli markaziy Kiosk boshqaruv ekrani.\n"
               "• `src/app/language/page.tsx` — Chet tillari moduli (Ibrat Farzandlari).\n"
               "• `src/app/career/page.tsx` — Zamonaviy kasblar moduli (Ustoz AI).\n"
               "• `src/app/migration/page.tsx` — Xavfsiz xorijiy migratsiya moduli.\n"
               "• `src/app/psychology/chat/page.tsx` — Psixologiya va 15 ta Hz shifobaxsh musiqa o'yini.\n"
               "• `src/app/history/page.tsx` — 5 allomalar virtual AI muzeyi.")

    # Save document
    output_filename = "Muqimiy_Aql_Markazi_Arxitekturasi_va_Avatar_Qollanmasi.docx"
    doc.save(output_filename)
    print(f"Document saved successfully as: {output_filename}")

if __name__ == '__main__':
    create_document()
